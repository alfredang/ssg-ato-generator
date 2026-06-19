from __future__ import annotations

import shutil
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
CLIENT = "Skills Union Pte Ltd"
OUT = ROOT / "ATO Supporting Document" / CLIENT
TODAY = date(2026, 6, 18)

SOURCE_FACTS = {
    "client": CLIENT,
    "source": "TC26-0606-20.pdf",
    "invoice_no": "TC26-0606-20",
    "invoice_date": "06/06/2026",
    "consultant": "Tertiary Infotech Academy Pte. Ltd.",
    "scope": "WSQ Courseware Submission & Development (1 Course) and WSQ ATO Organisation Registration (OR) Application support",
}


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E2F3"
TEXT = "1F1F1F"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "BFBFBF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)


def add_field_table(doc: Document, rows, widths=(2.0, 4.1), header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=0, cols=len(widths))
    set_table_width(table, widths)
    for idx, row_vals in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, val in enumerate(row_vals):
            cells[col_idx].text = str(val)
            cells[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cells[col_idx])
            if col_idx == 0 or idx == 0 and len(row_vals) > 2:
                set_cell_shading(cells[col_idx], header_fill)
                for p in cells[col_idx].paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def setup_doc(title: str, subtitle: str | None = None) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, "404040"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = f"{CLIENT} | SSG OR Supporting Document"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("666666")

    footer = section.footer.paragraphs[0]
    footer.text = "Prepared for SSG Organisation Registration submission"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("666666")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("666666")
    meta = doc.add_table(rows=1, cols=4)
    set_table_width(meta, [1.2, 2.0, 1.2, 1.7])
    row = meta.rows[0].cells
    vals = ["Client", CLIENT, "Prepared", TODAY.strftime("%d/%m/%Y")]
    for c, val in zip(row, vals):
        c.text = val
        set_cell_border(c, "D9D9D9")
        if val in ("Client", "Prepared"):
            set_cell_shading(c, LIGHT_BLUE)
            for r in c.paragraphs[0].runs:
                r.bold = True
    doc.add_paragraph()
    return doc


def add_source_note(doc: Document):
    doc.add_heading("Source Basis", level=2)
    add_field_table(
        doc,
        [
            ("Source file", SOURCE_FACTS["source"]),
            ("Invoice reference", SOURCE_FACTS["invoice_no"]),
            ("Invoice date", SOURCE_FACTS["invoice_date"]),
            ("Known scope", SOURCE_FACTS["scope"]),
            ("Information not present in source", "UEN, registered address, directors/shareholders, venue address, facilities photos, lease/rental evidence, historic training records."),
        ],
        widths=(1.7, 4.7),
    )


def acra_doc():
    doc = setup_doc("ACRA Business Profile Submission Cover", "Legal entity evidence pack")
    add_source_note(doc)
    doc.add_heading("Purpose", level=2)
    doc.add_paragraph(
        "This cover sheet is for the ACRA legal entity evidence required for SSG Organisation Registration. "
        "It should be uploaded together with the latest official ACRA Business Profile downloaded from BizFile."
    )
    doc.add_heading("Entity Details to Verify Against ACRA", level=2)
    add_field_table(
        doc,
        [
            ("Registered entity name", CLIENT),
            ("UEN", "[Client to insert from ACRA Business Profile]"),
            ("Entity type", "[Client to insert, e.g. Private Company Limited by Shares]"),
            ("Registration date", "[Client to insert]"),
            ("Registered office address", "[Client to insert from ACRA Business Profile]"),
            ("Principal activities", "[Client to insert SSIC activities shown in ACRA profile]"),
            ("Directors / officers", "[Client to insert names as shown in ACRA profile]"),
            ("Shareholders / members", "[Client to insert if disclosed in official profile]"),
        ],
        widths=(2.2, 4.2),
    )
    doc.add_heading("Upload Checklist", level=2)
    add_bullets(
        doc,
        [
            "Latest official ACRA Business Profile is attached, not older than 3 months at submission date.",
            "Entity name in TPGateway matches the legal entity name shown in ACRA.",
            "Registered office is not a P.O. Box.",
            "Key Person-In-Charge and Management Representatives match internal appointment records.",
            "Any differences between operating address and registered office are explained in the venue documents.",
        ],
    )
    path = OUT / "ACRA" / "Skills Union - ACRA Business Profile Cover.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def facilities_doc():
    doc = setup_doc("Facilities and Equipment Write-Up", "Premises, classroom and equipment evidence")
    doc.add_heading("Purpose", level=2)
    doc.add_paragraph(
        "This write-up supports the SSG OR premises requirement by describing the facilities and equipment available for training and assessment delivery."
    )
    doc.add_heading("Facilities Summary", level=2)
    add_field_table(
        doc,
        [
            ("Training provider", CLIENT),
            ("Training venue address", "[Client to insert venue address]"),
            ("Administrative office address", "[Client to insert office address]"),
            ("Room name / number", "[Client to insert]"),
            ("Maximum class size", "[Client to insert capacity]"),
            ("Venue ownership / lease status", "[Client to insert owned / leased / co-working / external venue]"),
            ("Primary course", "[Client to insert course title for OR and CA submission]"),
        ],
        widths=(2.2, 4.2),
    )
    doc.add_heading("Facilities and Equipment Available", level=2)
    rows = [
        ("Facility / equipment", "Availability / description", "Evidence to attach"),
        ("Training room", "[Client to describe seating layout, lighting, ventilation, accessibility]", "Photos of room from front, rear and side"),
        ("Trainer workstation", "[Client to describe laptop, display controls, internet access]", "Photo of trainer area"),
        ("Projection / display", "[Client to describe projector / TV / screen size]", "Photo of screen or projector"),
        ("Audio equipment", "[Client to describe speakers / microphone where applicable]", "Photo if applicable"),
        ("Learner seating", "[Client to describe desks, chairs, power access]", "Photo of learner seating"),
        ("Internet access", "[Client to describe Wi-Fi availability and backup]", "Screenshot / venue confirmation where available"),
        ("Assessment setup", "[Client to describe privacy, invigilation, device controls]", "Photos or procedure extract"),
        ("Safety and amenities", "[Client to describe exits, restrooms, pantry, first aid]", "Venue photos / floor plan if available"),
    ]
    add_field_table(doc, rows, widths=(1.55, 3.1, 1.75), header_fill=LIGHT_BLUE)
    doc.add_heading("Photo Log", level=2)
    rows = [("Photo ref", "Description", "File name / status")]
    for i in range(1, 7):
        rows.append((f"Photo {i}", "[Client to insert photo description]", "[Attach image file]"))
    add_field_table(doc, rows, widths=(1.1, 3.6, 1.7), header_fill=LIGHT_BLUE)
    path = OUT / "Facilities" / "Skills Union - Facilities and Equipment Write-Up.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def training_venue_doc():
    doc = setup_doc("Training Venue Declaration and Inspection Checklist", "Venue readiness for SSG-funded training delivery")
    doc.add_heading("Venue Declaration", level=2)
    add_field_table(
        doc,
        [
            ("Training provider", CLIENT),
            ("Venue address", "[Client to insert full venue address]"),
            ("Administrative office address", "[Client to insert full office address]"),
            ("Venue contact person", "[Client to insert name, designation, phone and email]"),
            ("Period of access", "[Client to insert lease / rental period]"),
            ("Classroom capacity", "[Client to insert maximum learners]"),
            ("Assessment capacity", "[Client to insert maximum candidates]"),
        ],
        widths=(2.2, 4.2),
    )
    doc.add_heading("Inspection Checklist", level=2)
    rows = [
        ("Area", "Requirement", "Status", "Remarks / evidence"),
        ("Location", "Venue address is a physical location and not a P.O. Box.", "[Pass / Pending]", "[Client to insert]"),
        ("Access", "Learners can locate and access the training room safely.", "[Pass / Pending]", "[Client to insert]"),
        ("Capacity", "Room capacity supports intended class size with suitable spacing.", "[Pass / Pending]", "[Client to insert]"),
        ("Equipment", "Projector/display, trainer device and internet are available.", "[Pass / Pending]", "[Client to insert]"),
        ("Assessment", "Assessment setup supports identity checks and invigilation.", "[Pass / Pending]", "[Client to insert]"),
        ("Records", "Lease, rental invoice or venue booking proof is available.", "[Pass / Pending]", "[Client to insert]"),
        ("Safety", "Emergency exits, restrooms and basic amenities are accessible.", "[Pass / Pending]", "[Client to insert]"),
    ]
    add_field_table(doc, rows, widths=(1.2, 2.8, 1.1, 1.3), header_fill=LIGHT_BLUE)
    doc.add_heading("Documents to Attach", level=2)
    add_bullets(
        doc,
        [
            "Latest lease agreement, rental invoice or venue booking confirmation.",
            "Photos of training room and assessment setup.",
            "Floor plan or wayfinding details, where available.",
            "Any external venue agreement showing right of use for training delivery.",
        ],
    )
    doc.add_heading("Approval", level=2)
    add_field_table(
        doc,
        [
            ("Prepared by", "[Name / designation]"),
            ("Reviewed by", "[Name / designation]"),
            ("Date", "[DD/MM/YYYY]"),
            ("Signature", "[Signature or e-signature]"),
        ],
        widths=(2.0, 4.4),
    )
    path = OUT / "Training Venue" / "Skills Union - Training Venue Declaration and Inspection Checklist.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def systems_doc():
    doc = setup_doc("Systems and Capabilities Manual", "Policies and operations manual for SSG Organisation Registration")
    add_source_note(doc)
    doc.add_heading("Document Control", level=2)
    add_field_table(
        doc,
        [
            ("Document owner", "[Client to insert Key Person-In-Charge]"),
            ("Version", "1.0"),
            ("Effective date", "[Client to insert]"),
            ("Review cycle", "Annual or upon material process change"),
            ("Prepared for", "SSG Organisation Registration application"),
        ],
        widths=(2.0, 4.4),
    )
    doc.add_heading("1. Governance and Organisation", level=1)
    doc.add_paragraph(
        f"{CLIENT} will maintain clear accountability for SSG-funded course administration, learner support, assessment integrity, finance, records management and continuous improvement."
    )
    add_field_table(
        doc,
        [
            ("Role", "Key responsibilities"),
            ("Key Person-In-Charge", "Overall compliance with SSG terms, approval of course information, oversight of audit responses and management reviews."),
            ("Management Representative", "Day-to-day monitoring of course administration, learner communications, trainer deployment and document control."),
            ("Course Administrator", "Enrolment, course confirmation, attendance, assessment records, claims support and trainee communications."),
            ("Trainer / Assessor", "Lesson delivery, assessment conduct, learner support, feedback collection and post-course review input."),
            ("Finance Representative", "Fee collection, invoice records, refund processing, financial reconciliation and supporting evidence retention."),
        ],
        widths=(2.0, 4.4),
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("2. Course Information and Learner Communication", level=1)
    doc.add_paragraph(
        "Published course information must be accurate, approved and consistent across website, brochure, registration form and confirmation emails."
    )
    add_numbered(
        doc,
        [
            "Course Administrator drafts or updates course information using the approved course factsheet.",
            "Management Representative checks course title, duration, mode, objectives, fees, funding information, trainers, facilities and completion criteria.",
            "Key Person-In-Charge approves release before publication or learner circulation.",
            "All approved versions are stored in a controlled folder with version date and owner.",
        ],
    )
    doc.add_heading("Required Course Disclosures", level=2)
    add_bullets(
        doc,
        [
            "Course title, training duration, fees and funding validity period.",
            "Modes of training and assessment.",
            "Course objectives and intended learning outcomes.",
            "Senior management staff and trainer names.",
            "Organisation structure and facilities/equipment used for training.",
            "No SSG or SkillsFuture logos are used in marketing materials unless expressly permitted.",
        ],
    )

    doc.add_heading("3. Registration, Course Confirmation and Advisory", level=1)
    doc.add_paragraph(
        "Learner registration records must show the registered course title, commencement and end dates, funding to be applied and balance fees payable by learners."
    )
    add_field_table(
        doc,
        [
            ("Process step", "Control"),
            ("Pre-course advisory", "Confirm course fit, prerequisites, funding conditions, attendance and assessment requirements before enrolment."),
            ("Registration", "Collect learner particulars, company billing details where applicable, declaration and consent."),
            ("Course confirmation", "Issue confirmation with course schedule, venue/login details, trainer, materials, attendance criteria and support contact."),
            ("Changes and withdrawals", "Record requests, communicate options and apply refund or deferment rules consistently."),
            ("Appeals", "Log appeal, appoint reviewer, respond with outcome and keep evidence of decision."),
        ],
        widths=(2.0, 4.4),
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("4. Attendance, Assessment and Records", level=1)
    add_bullets(
        doc,
        [
            "Attendance is taken using the approved e-attendance method required for classroom-facilitated training or synchronous e-learning.",
            "Trainer verifies learner identity before attendance and assessment activities.",
            "Assessment instruments, answer scripts, assessor decisions and moderation records are retained according to the internal retention schedule.",
            "Records are stored in access-controlled folders with backup and version control.",
            "Assessment results are reviewed before submission to TPGateway or other required systems.",
        ],
    )
    doc.add_heading("Records Register", level=2)
    add_field_table(
        doc,
        [
            ("Record type", "Retention / control"),
            ("Learner registration form", "Retain per SSG and internal data retention policy; restrict access to authorised staff."),
            ("Attendance record", "Retain source record and submission evidence."),
            ("Assessment evidence", "Retain assessment plan, instruments, learner evidence, marking records and moderation evidence."),
            ("Invoices and payment records", "Retain issued invoice, receipt, refund record and reconciliation evidence."),
            ("Feedback and complaints", "Retain feedback forms, analysis, action plan and closure evidence."),
        ],
        widths=(2.1, 4.3),
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("5. Trainer Management and Course Quality", level=1)
    add_numbered(
        doc,
        [
            "Verify trainer qualifications, industry experience and ability to deliver the course before appointment.",
            "Conduct briefing on course objectives, lesson plan, assessment requirements, attendance controls and learner support procedures.",
            "Collect learner feedback and trainer reflection after each run.",
            "Review attendance, completion, assessment outcomes and feedback trends during management review.",
            "Implement corrective actions where learner outcomes, feedback or audit observations indicate gaps.",
        ],
    )

    doc.add_heading("6. Finance, Funding and Refund Controls", level=1)
    add_bullets(
        doc,
        [
            "Course fees, SkillsFuture Credit and other SSG funding amounts are communicated before registration.",
            "Invoices and receipts are matched against registration records.",
            "Refunds, withdrawals and deferments are approved according to published policy.",
            "Finance records are reconciled and retained as supporting evidence for audit or clarification.",
            "The organisation will maintain a PayNow account for TPGateway transactions where required.",
        ],
    )

    doc.add_heading("7. Feedback, Complaints, Requests and Appeals", level=1)
    add_field_table(
        doc,
        [
            ("Channel", "Handling standard"),
            ("General enquiry", "Acknowledge within 2 working days and route to the responsible staff member."),
            ("Feedback", "Log feedback, classify theme, review for improvement action and close with documented outcome."),
            ("Complaint", "Escalate to Management Representative, investigate evidence, reply with outcome and corrective action."),
            ("Appeal", "Assign an independent reviewer where possible and document final decision."),
        ],
        widths=(1.8, 4.6),
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("8. Data Protection and Access Control", level=1)
    add_bullets(
        doc,
        [
            "Personal data is collected only for course administration, regulatory submission and learner support purposes.",
            "Access to learner records is restricted to authorised staff and reviewed when staff roles change.",
            "Files are stored in controlled folders with clear naming conventions and backups.",
            "Data incidents are escalated to management for containment, assessment and notification where required.",
        ],
    )

    doc.add_heading("9. Management Review and Continuous Improvement", level=1)
    doc.add_paragraph(
        "Management will review operational performance at least annually and after significant course runs, audit findings or complaints."
    )
    add_field_table(
        doc,
        [
            ("Review input", "Example measures"),
            ("Course administration", "Enrolment accuracy, confirmation timeliness, attendance submission accuracy."),
            ("Learner outcomes", "Completion rates, assessment outcomes, learner support cases."),
            ("Trainer performance", "Learner feedback, trainer observation, assessment quality."),
            ("Compliance", "Audit findings, SSG clarification trends, document control status."),
            ("Improvement actions", "Owner, due date, evidence of closure and effectiveness review."),
        ],
        widths=(2.0, 4.4),
        header_fill=LIGHT_BLUE,
    )

    doc.add_heading("Appendix A: Client Completion Items", level=1)
    add_bullets(
        doc,
        [
            "Insert named appointment holders and organisation chart.",
            "Insert actual course title, registration form, brochure or website mock-up.",
            "Attach venue evidence, facilities photos and training records.",
            "Confirm e-attendance method, TMS/LMS platform and document storage location.",
            "Insert actual refund policy, appeal form and feedback form if already in use.",
        ],
    )
    path = OUT / "System and Capabilities" / "Skills Union - Systems and Capabilities Manual.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def training_records_xlsx():
    path = OUT / "Training Records" / "Skills Union - Training Track Records Template.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Training Track Records"
    ws.merge_cells("A1:H1")
    ws["A1"] = "WRITE UP FOR TRAINING AND EDUCATION ACTIVITIES CONDUCTED IN PRECEDING ONE YEAR"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor=BLUE)
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A3:H3")
    ws["A3"] = "Client: Skills Union Pte Ltd | Prepared for SSG Organisation Registration | Source: TC26-0606-20.pdf"
    ws["A3"].font = Font(italic=True, color="666666")
    notes = [
        "All fields are mandatory unless marked optional.",
        "Create a separate entry for each course run conducted in the preceding one year.",
        "Dates should be entered in DD/MM/YYYY format.",
        "Attach evidence such as invoices, attendance records, confirmation emails, lesson plans, assessment plans and learner feedback.",
        "The invoice source does not include Skills Union's historic training records; the rows below are intentionally blank for client completion.",
    ]
    for i, note in enumerate(notes, start=5):
        ws[f"A{i}"] = i - 4
        ws[f"B{i}"] = note
        ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=8)
    header_row = 12
    headers = [
        "S/N",
        "Course title",
        "Duration (No. of Hours)",
        "Course commencement date (DD/MM/YYYY)",
        "Course end date (DD/MM/YYYY)",
        "Age profile of trainees",
        "Proof of Training Conducted submitted? (Yes/No)",
        "Remarks / Evidence references",
    ]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(header_row, col, header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in range(header_row + 1, header_row + 31):
        ws.cell(row, 1, row - header_row)
        for col in range(2, 9):
            ws.cell(row, col, "")
    widths = [8, 36, 16, 22, 22, 22, 24, 36]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = "A13"
    ws.auto_filter.ref = f"A{header_row}:H{header_row + 30}"
    thin = Side(style="thin", color="BFBFBF")
    for row in ws.iter_rows(min_row=1, max_row=header_row + 30, min_col=1, max_col=8):
        for cell in row:
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    for row in range(header_row + 1, header_row + 31):
        ws.row_dimensions[row].height = 42
    ws.row_dimensions[header_row].height = 48

    ev = wb.create_sheet("Evidence Index")
    ev_headers = ["Evidence ref", "Evidence type", "Related course run", "File name", "Notes"]
    for col, header in enumerate(ev_headers, start=1):
        c = ev.cell(1, col, header)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=BLUE)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    examples = [
        ("EV-001", "Invoice", "[S/N]", "[Attach file name]", "Invoice billed to trainee or corporate client."),
        ("EV-002", "Attendance record", "[S/N]", "[Attach file name]", "Signed attendance or e-attendance export."),
        ("EV-003", "Confirmation email", "[S/N]", "[Attach file name]", "Course registration confirmation and essential course information."),
        ("EV-004", "Lesson / assessment plan", "[S/N]", "[Attach file name]", "Shows learning objectives, structured activities and assessment approach."),
        ("EV-005", "Feedback / evaluation", "[S/N]", "[Attach file name]", "Shows learner feedback and course evaluation analysis."),
    ]
    for row, values in enumerate(examples, start=2):
        for col, val in enumerate(values, start=1):
            ev.cell(row, col, val)
    for idx, width in enumerate([14, 22, 20, 34, 48], start=1):
        ev.column_dimensions[get_column_letter(idx)].width = width
    for row in ev.iter_rows(min_row=1, max_row=20, min_col=1, max_col=5):
        for cell in row:
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    ev.freeze_panes = "A2"
    wb.save(path)
    return path


def training_records_cover():
    doc = setup_doc("Training Records Evidence Pack Cover", "Training track records and supporting evidence index")
    doc.add_heading("Purpose", level=2)
    doc.add_paragraph(
        "This cover sheet accompanies the training track record workbook and evidence files required for SSG OR Stage 1."
    )
    doc.add_heading("SSG Track Record Requirements Reflected in This Pack", level=2)
    add_bullets(
        doc,
        [
            "At least one-year training and education track record.",
            "Regular conduct of training, at least once in each quarter of the preceding year.",
            "Course title, duration, course commencement and end dates, and age profile of trainees for each run.",
            "Evidence of actual training activities, such as invoices, attendance records, course confirmation communications and recordings for online training.",
            "Evidence of course objectives, structured lesson plans, assessment plans, industry engagement and learner feedback analysis.",
        ],
    )
    doc.add_heading("Client Completion Note", level=2)
    doc.add_paragraph(
        "The source invoice does not contain Skills Union's historic training runs. Complete the workbook with actual completed training activities and attach source evidence before submission."
    )
    path = OUT / "Training Records" / "Skills Union - Training Records Evidence Pack Cover.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def render_docx(docx_paths):
    renderer = Path("/Users/alfredang/.codex/plugins/cache/openai-primary-runtime/documents/26.614.11602/skills/documents/render_docx.py")
    python = Path("/Users/alfredang/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3")
    rendered = []
    for docx_path in docx_paths:
        out_dir = ROOT / "tmp" / "rendered" / docx_path.stem
        out_dir.mkdir(parents=True, exist_ok=True)
        subprocess.run([str(python), str(renderer), str(docx_path), "--output_dir", str(out_dir), "--emit_pdf"], check=True)
        rendered.append(out_dir)
    return rendered


def convert_systems_pdf(docx_path: Path):
    out_dir = docx_path.parent
    soffice = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if soffice.exists():
        subprocess.run([str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)], check=True)
        return out_dir / (docx_path.stem + ".pdf")
    return None


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    docx_paths = [
        acra_doc(),
        facilities_doc(),
        training_records_cover(),
        systems_doc(),
        training_venue_doc(),
    ]
    xlsx_path = training_records_xlsx()
    source_copy = OUT / "00 Source" / "TC26-0606-20.pdf"
    source_copy.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(ROOT / "TC26-0606-20.pdf", source_copy)
    pdf_path = convert_systems_pdf(docx_paths[3])
    rendered = render_docx(docx_paths)
    print("Generated:")
    for p in docx_paths + [xlsx_path, source_copy]:
        print(p)
    if pdf_path:
        print(pdf_path)
    print("Rendered QA:")
    for p in rendered:
        print(p)


if __name__ == "__main__":
    main()
