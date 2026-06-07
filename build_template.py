"""Author the standard, placeholder-only SOP template (sop_template.docx).

The template contains a cover page, a version-control record, an auto-updating
Table of Contents, all SOP sections (with diagram markers for the recreated
flowcharts/org charts), and a company-name footer. It holds ZERO real company
or personal details -- only {{PLACEHOLDERS}} and [[DIAGRAM:key]] markers.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

NAVY = RGBColor(0x1F, 0x3A, 0x5F)

# ---- diagram marker helper ----
def DIAG(key):
    return f"[[DIAGRAM:{key}]]"


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def centered(doc, text, size=12, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def diagram(doc, key, caption=None):
    p = doc.add_paragraph(DIAG(key))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Update this field to build the Table of Contents."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, instr, f2, t, f3):
        run._r.append(el)


def enable_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def set_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("{{COMPANY_NAME}}  |  Internal Document")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def build():
    doc = Document()

    # base styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for lvl, sz in ((1, 16), (2, 13), (3, 12)):
        st = doc.styles[f"Heading {lvl}"]
        st.font.color.rgb = NAVY
        st.font.size = Pt(sz)

    # ---------------- COVER PAGE ----------------
    for _ in range(2):
        doc.add_paragraph()
    # logo marker (app inserts image here if logo toggle on, else removes line)
    logo_p = doc.add_paragraph("{{LOGO}}")
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    centered(doc, "Training Systems and Capabilities", size=28, bold=True, color=NAVY, space_after=4)
    centered(doc, "Standard Operating Procedure", size=16, italic=True, space_after=24)
    centered(doc, "{{COMPANY_NAME}}", size=20, bold=True, space_after=2)
    centered(doc, "UEN: {{UEN}}", size=11, space_after=2)
    centered(doc, "{{ADDRESS}}", size=11, space_after=2)
    centered(doc, "{{WEBSITE}}", size=11, space_after=18)
    centered(doc, "Support Contact", size=12, bold=True, space_after=2)
    centered(doc, "{{CONTACT_NAME}}", size=11, space_after=1)
    centered(doc, "Tel: {{CONTACT_TEL}}", size=11, space_after=1)
    centered(doc, "Email: {{CONTACT_EMAIL}}", size=11, space_after=1)
    add_page_break(doc)

    # ---------------- VERSION CONTROL ----------------
    doc.add_heading("Version Control Record", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Version", "Effective Date", "Change History", "Updated By"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for ver in (("1.0", "{{DATE}}", "1st Release", "{{OWNER_NAME}}"),):
        cells = tbl.add_row().cells
        for i, v in enumerate(ver):
            cells[i].text = v
    body(doc, "")
    add_page_break(doc)

    # ---------------- TABLE OF CONTENTS ----------------
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    add_page_break(doc)

    # ---------------- ORG CHART (management) ----------------
    doc.add_heading("Organisation Chart of Management Team", level=1)
    body(doc, "The management team of {{COMPANY_NAME}} is led by {{OWNER_NAME}}. "
              "The organisation chart below sets out the reporting lines across the "
              "company's functional departments.")
    diagram(doc, "org_management")
    add_page_break(doc)

    # ---------------- A. COURSE ADMINISTRATION ----------------
    doc.add_heading("A. Course Administration", level=1)

    doc.add_heading("A.1  Roles and Responsibilities", level=2)
    body(doc, "{{COMPANY_NAME}} defines clear roles and responsibilities for the "
              "delivery and administration of its training programmes. The Training "
              "Manager and Management Representative oversees overall training quality "
              "and compliance with SkillsFuture Singapore (SSG) requirements, supported "
              "by the finance, administration and training functions.")

    doc.add_heading("A.2  Organisation Chart of Training Team", level=2)
    body(doc, "The training team structure of {{COMPANY_NAME}} is shown below.")
    diagram(doc, "org_training")

    doc.add_heading("A.3  Administrative Support and Processes", level=2)
    body(doc, "Enquiries and registrations are handled through a standardised process "
              "to ensure prospective trainees receive timely advisory and accurate "
              "course information. The end-to-end pre-course enquiry and registration "
              "flow is shown below.")
    diagram(doc, "enquiry")
    body(doc, "Upon successful enrolment, the administration team confirms the funding "
              "route with the trainee, whether self-sponsored or employer-sponsored, and "
              "guides them on the use of SkillsFuture credits where applicable.")
    diagram(doc, "funding")
    body(doc, "Before each course run, {{COMPANY_NAME}} confirms the run and sends "
              "reminders to enrolled trainees, or notifies them of any cancellation and "
              "processes refunds accordingly.")
    diagram(doc, "course_confirmation")

    doc.add_heading("A.4  Marketing Guidelines by SSG", level=2)
    body(doc, "All marketing collateral published by {{COMPANY_NAME}} complies with the "
              "SSG marketing guidelines, including accurate representation of course fees, "
              "funding, trainer credentials and course outcomes.")

    doc.add_heading("A.5  Synchronous E-Training Delivery System", level=2)
    body(doc, "For synchronous e-learning conducted over an approved video-conferencing "
              "platform, {{COMPANY_NAME}} applies the following controls:")
    bullets(doc, [
        "Authenticate learners' identities (e.g. pre-course identity verification).",
        "Verify attendance at the start, mid-point and end of each session.",
        "Record sessions where required, with the trainer and Training Co-ordinator "
        "holding the Zoom (or equivalent) licence and host controls.",
        "Capture screenshots / attendance evidence to support attendance submission.",
    ])

    doc.add_heading("A.6  Learner Support System", level=2)
    body(doc, "{{COMPANY_NAME}} provides learners with technical and learning support "
              "before and during the course, including guidance on platform access, "
              "course materials and assessment requirements. Learners may contact "
              "{{CONTACT_NAME}} at {{CONTACT_EMAIL}} or {{CONTACT_TEL}} for assistance.")

    doc.add_heading("A.7  Attendance Submission", level=2)
    body(doc, "Attendance is recorded using the mandated e-attendance system and "
              "submitted to SSG within the stipulated timelines, with supporting "
              "verification records retained for audit.")

    doc.add_heading("A.8  Pre-course and Post-course Advisory Service", level=2)
    body(doc, "A learning consultant provides pre-course advisory on course suitability, "
              "fees, funding and eligibility, and post-course advisory on progression "
              "pathways and further learning opportunities.")

    doc.add_heading("A.9  Assessment Appeal", level=2)
    body(doc, "Trainees may appeal an assessment outcome in writing within the stipulated "
              "period. {{COMPANY_NAME}} reviews each appeal objectively and informs the "
              "trainee of the outcome.")

    doc.add_heading("A.10  Refund Process", level=2)
    body(doc, "Refund requests are received in writing and assessed against the refund "
              "policy and course start date, then routed for approval as shown below.")
    diagram(doc, "refund")

    doc.add_heading("A.11  Documentation, Record Keeping and Review of SOPs", level=2)
    body(doc, "{{COMPANY_NAME}} maintains records of registrations, attendance, "
              "assessments and evaluations in line with SSG retention requirements. These "
              "Standard Operating Procedures are reviewed at least annually, or when there "
              "is a material change in regulations or operations, by {{OWNER_NAME}}.")
    add_page_break(doc)

    # ---------------- B. TRAINING QUALITY ----------------
    doc.add_heading("B. Training Quality", level=1)

    doc.add_heading("B.12  Trainer Management and Processes", level=2)
    body(doc, "{{COMPANY_NAME}} recruits, on-boards and evaluates trainers and assessors "
              "against defined competency criteria. Trainer performance is monitored "
              "through course evaluations; where a performance issue is identified, the "
              "following process applies.")
    diagram(doc, "trainer_performance")

    doc.add_heading("B.13  Content Design and Development and Processes", level=2)
    body(doc, "Course content is designed and developed with reference to industry and "
              "sector needs, learning outcomes and SSG curriculum requirements. Materials "
              "undergo pre-delivery and post-delivery evaluation to ensure relevance and "
              "quality.")
    add_page_break(doc)

    # ---------------- C. TRAINING OUTCOMES ----------------
    doc.add_heading("C. Training Outcomes", level=1)

    doc.add_heading("C.14  Outcome Evaluation, Data Collection and Analysis", level=2)
    body(doc, "{{COMPANY_NAME}} collects and analyses training outcome data across "
              "reaction (L1), learning (L2), behaviour (L3) and results (L4) levels to "
              "evaluate the effectiveness of its programmes.")

    doc.add_heading("C.15  Recommendations and Follow-Up", level=2)
    body(doc, "Findings from outcome evaluation are reviewed by management, and "
              "recommendations are implemented and tracked to drive continuous improvement "
              "of {{COMPANY_NAME}}'s training quality and service.")
    add_page_break(doc)

    # ---------------- D. ANNEX SUMMARY ----------------
    doc.add_heading("D. Annex Summary", level=1)
    body(doc, "The following annex templates support these SOPs and are maintained by "
              "{{COMPANY_NAME}}:")
    bullets(doc, [
        "Annex 1  Course Registration Form",
        "Annex 2  Course Confirmation & Agreement Template",
        "Annex 3  Attendance Sheet Template",
        "Annex 4  Refund Process",
        "Annex 5  Review of Assessment Appeal Case Form",
        "Annex 6  Assessment Appeal Procedure",
        "Annex 7  Interview / Selection Evaluation Form for Trainer / Assessor",
        "Annex 8  Trainer / Assessor & Developer Competency Checklists",
        "Annex 9  Course Evaluation Forms (L1-L4) and Consolidated Reports",
        "Annex 10 Minutes of Meeting Template",
    ])

    # footer + field update
    set_footer(doc)
    enable_update_fields(doc)

    doc.save("sop_template.docx")
    print("wrote sop_template.docx")


if __name__ == "__main__":
    build()
