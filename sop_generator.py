"""Core SOP generation: fill the placeholder template for a given company.

- Replaces {{PLACEHOLDER}} tokens (body, tables, headers, footers).
- Regenerates diagrams with the company's owner name and inserts them at
  [[DIAGRAM:key]] markers.
- Handles the {{LOGO}} marker via a show/hide toggle.
"""
import io
import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

import flowcharts

TEMPLATE = Path(__file__).parent / "sop_template.docx"

PLACEHOLDER_KEYS = [
    "COMPANY_NAME", "OWNER_NAME", "UEN", "ADDRESS", "WEBSITE",
    "CONTACT_NAME", "CONTACT_TEL", "CONTACT_EMAIL", "DATE",
]

_DIAG_RE = re.compile(r"\[\[DIAGRAM:([a-z_]+)\]\]")


def _replace_in_paragraph(p, mapping):
    """Replace {{KEY}} tokens within a paragraph, preserving run formatting."""
    full = "".join(r.text for r in p.runs)
    if "{{" not in full:
        return
    new = full
    for k, v in mapping.items():
        new = new.replace("{{" + k + "}}", v)
    if new == full:
        return
    # write everything into the first run, clear the rest (keeps first run format)
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""


def _iter_paragraphs(doc):
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer):
            yield from hf.paragraphs


def _fit_box(img_path, max_w=6.3, max_h=8.2):
    """Return (width_in, height_in) fitting the image into a max box, keeping aspect."""
    with Image.open(img_path) as im:
        w, h = im.size
    ar = w / h
    width = max_w
    height = width / ar
    if height > max_h:
        height = max_h
        width = height * ar
    return width, height


def _insert_image(paragraph, img_path):
    for r in list(paragraph.runs):
        r.text = ""
    w, h = _fit_box(img_path)
    run = paragraph.add_run()
    run.add_picture(img_path, width=Inches(w), height=Inches(h))


def generate(info, logo_bytes=None, show_logo=True, template_path=TEMPLATE):
    """Generate a filled SOP. `info` keys -> PLACEHOLDER_KEYS. Returns BytesIO."""
    mapping = {k: (info.get(k.lower(), "") or f"[{k}]") for k in PLACEHOLDER_KEYS}

    doc = Document(str(template_path))

    # 1) regenerate diagrams with the owner name
    owner = info.get("owner_name") or "[Owner / Managing Director]"
    diag_dir = tempfile.mkdtemp(prefix="sopdiag_")
    diagrams = flowcharts.build_all(diag_dir, owner=owner)

    # 2) text placeholders
    for p in _iter_paragraphs(doc):
        _replace_in_paragraph(p, mapping)

    # 3) logo marker + diagram markers (body paragraphs)
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text == "{{LOGO}}":
            if show_logo and logo_bytes:
                tmp = os.path.join(diag_dir, "logo.png")
                with open(tmp, "wb") as f:
                    f.write(logo_bytes)
                try:
                    _insert_image(p, tmp)
                except Exception:
                    p.text = ""
            else:
                p.text = ""
            continue
        m = _DIAG_RE.search(text)
        if m and m.group(1) in diagrams:
            _insert_image(p, diagrams[m.group(1)])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def safe_filename(company_name):
    base = re.sub(r"[^A-Za-z0-9_-]+", "_", (company_name or "Company").strip()).strip("_")
    return f"SOP_{base or 'Company'}.docx"
