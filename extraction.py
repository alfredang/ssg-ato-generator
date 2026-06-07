"""Extract company details from uploaded supporting documents.

Best-effort parsing of ACRA / BizFile business profiles and similar documents
(PDF or DOCX) to pre-fill the SOP form: company name, UEN, registered address,
contact email and telephone. Returns whatever it can find; the user can always
review and edit before generating.
"""
import io
import re
import subprocess
import tempfile

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
TEL_RE = re.compile(r"(?:\+65[\s-]?)?[689]\d{3}[\s-]?\d{4}")
UEN_RE = re.compile(r"\b(\d{8,10}[A-Z]|[ST]\d{2}[A-Z]{2}\d{4}[A-Z])\b")


def _pdf_text(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(data)
        path = f.name
    try:
        out = subprocess.run(["pdftotext", "-layout", path, "-"],
                             capture_output=True, text=True, timeout=60)
        text = out.stdout or ""
    except Exception:
        text = ""
    if len(text.strip()) < 20:
        # fallback to pypdf
        try:
            from pypdf import PdfReader
            r = PdfReader(path)
            text = "\n".join((p.extract_text() or "") for p in r.pages)
        except Exception:
            pass
    return text


def _docx_text(data: bytes) -> str:
    try:
        from docx import Document
        d = Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    parts.append(c.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_fields(text: str) -> dict:
    info = {}
    # Company name: "Business Profile (Company) of XXX (UEN)" or "Name of Company : XXX"
    m = re.search(r"Business Profile.*?of\s+(.+?)\s*\(\s*\d", text, re.I | re.S)
    if not m:
        m = re.search(r"Name of Company\s*[:\-]?\s*(.+)", text, re.I)
    if m:
        info["company_name"] = re.sub(r"\s+", " ", m.group(1)).strip()

    m = re.search(r"UEN\s*[:\-]?\s*([0-9A-Z]+)", text, re.I)
    if m:
        info["uen"] = m.group(1).strip()
    elif UEN_RE.search(text):
        info["uen"] = UEN_RE.search(text).group(1)

    m = re.search(r"Registered Office Address\s*[:\-]?\s*(.+)", text, re.I)
    if m:
        info["address"] = re.sub(r"\s{2,}", " ", m.group(1)).strip()

    # Director / owner
    m = re.search(r"(?:Name)\s*\n?\s*Address.*?\n\s*([A-Z][A-Z\s./]+?)\s+[STFG]\d{6,}",
                  text, re.S)
    m2 = re.search(r"([A-Z][A-Za-z\s./]+?)\s+.*?DIRECTOR", text)
    if m2:
        info["owner_name"] = re.sub(r"\s+", " ", m2.group(1)).strip().title()

    em = EMAIL_RE.search(text)
    if em:
        info["contact_email"] = em.group(0)
    tel = TEL_RE.search(text)
    if tel:
        info["contact_tel"] = tel.group(0)
    return info


def extract_company_info(uploaded_files) -> tuple[dict, list]:
    """uploaded_files: list of Streamlit UploadedFile. Returns (info, notes)."""
    info, notes = {}, []
    for uf in uploaded_files or []:
        name = getattr(uf, "name", "file")
        data = uf.read()
        try:
            uf.seek(0)
        except Exception:
            pass
        low = name.lower()
        if low.endswith(".pdf"):
            text = _pdf_text(data)
        elif low.endswith((".docx",)):
            text = _docx_text(data)
        else:
            notes.append(f"Skipped '{name}' (unsupported type for extraction).")
            continue
        if len(text.strip()) < 20:
            notes.append(f"'{name}': no machine-readable text found "
                         f"(likely a scanned image — please enter details manually).")
            continue
        found = _extract_fields(text)
        for k, v in found.items():
            info.setdefault(k, v)
        if found:
            notes.append(f"'{name}': extracted {', '.join(found.keys())}.")
        else:
            notes.append(f"'{name}': no recognisable fields found.")
    return info, notes
